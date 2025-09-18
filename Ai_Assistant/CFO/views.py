from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.views.decorators.cache import never_cache
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse
from django.template.loader import render_to_string
import pandas as pd
import os
import json
import re
import csv
import openai
from io import BytesIO
from openpyxl import Workbook
from reportlab.pdfgen import canvas
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import base64
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import SetPasswordForm
from django.contrib import messages
from django.shortcuts import render, redirect
from .forms import ProfileForm, PreferencesForm
from .models import Profile

# PDF/DOCX
import pdfplumber
from docx import Document

import io
import tempfile
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from django.contrib.auth.decorators import login_required
import matplotlib
matplotlib.use('Agg')  # <- Use non-GUI backend
import matplotlib.pyplot as plt

from .models import FinancialReport

import base64

# OCR
from pdf2image import convert_from_path
import pytesseract

from .forms import (
    FinancialReportForm, SignUpForm, LoginForm, ScenarioForm, 
    ReportFilterForm, GoalForm, ProfileForm, PreferencesForm
)
from .models import FinancialReport, FinancialGoal, Contact, Profile, UserPreferences


# ---------------- Helper Functions ----------------
def clean_numeric_series(series):
    return pd.to_numeric(series.replace(r'[^0-9.\-]', '', regex=True), errors='coerce').fillna(0)


def find_col_by_keywords(columns, keywords):
    for key in keywords:
        for col in columns:
            if key.lower() in col.lower():
                return col
    return None


def normalize_financial_df(df):
    """Normalize any uploaded financial data into Month, Account, Amount format."""
    if df is None or df.empty:
        return pd.DataFrame(columns=["Month", "Account", "Amount"])
    df.columns = df.columns.str.strip()

    if "Amount" in df.columns and "Account" in df.columns:
        df["Amount"] = clean_numeric_series(df["Amount"])
        if "Month" not in df.columns:
            df["Month"] = "Full Year"
        return df[["Month", "Account", "Amount"]]

    month_col = find_col_by_keywords(df.columns, ["month", "date", "period"])
    if month_col:
        df[month_col] = pd.to_datetime(df[month_col], errors="coerce")
        df["Month"] = df[month_col].dt.to_period("M").astype(str)
    else:
        df["Month"] = "Full Year"

    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if not numeric_cols:
        for col in df.columns:
            df[col] = pd.to_numeric(df[col].replace(r"[^0-9.\-]", "", regex=True), errors="coerce")
        numeric_cols = df.select_dtypes(include="number").columns.tolist()

    if not numeric_cols:
        cols = df.columns.tolist()
        df["Account"] = df[cols[0]] if len(cols) > 0 else "Unknown"
        df["Amount"] = pd.to_numeric(df[cols[1]], errors="coerce").fillna(0) if len(cols) > 1 else 0

    safe_cols = {col: f"{col}_val" for col in numeric_cols}
    df = df.rename(columns=safe_cols)

    melted = df.melt(
        id_vars=["Month"],
        value_vars=list(safe_cols.values()),
        var_name="Account",
        value_name="Amount"
    )

    melted["Account"] = melted["Account"].str.replace("_val", "", regex=False)
    melted["Amount"] = melted["Amount"].fillna(0)

    if melted['Month'].nunique() == 1 and melted['Month'].iloc[0] == "Full Year":
        melted['Month'] = [f"Month {i+1}" for i in range(len(melted))]

    return melted[["Month", "Account", "Amount"]]



# Add this to your views.py for AI assistant functionality

@login_required(login_url='CFO:login')
def ai_assistant(request):
    """AI-powered financial assistant view"""
    if request.method == 'POST':
        user_query = request.POST.get('query', '')
        
        # Get user's latest financial data for context
        latest_report = FinancialReport.objects.filter(user=request.user).order_by('-uploaded_at').first()
        
        context_data = {}
        if latest_report:
            context_data = {
                'revenue': latest_report.total_revenue or 0,
                'profit': latest_report.total_profit or 0,
                'expense': latest_report.total_expense or 0,
                'cash_balance': latest_report.cash_balance or 0,
                'profit_margin': latest_report.profit_margin or 0,
                'runway_months': latest_report.runway_months or 0
            }
        
        # Generate AI response
        ai_response = generate_financial_advice(user_query, context_data)
        
        return JsonResponse({
            'response': ai_response,
            'query': user_query
        })
    
    return render(request, 'CFO/ai_assistant.html')

def generate_financial_advice(query, financial_data):
    """Generate AI-powered financial advice using context"""
    
    # Create context prompt
    context_prompt = f"""
    You are a CFO AI Assistant. Here's the user's current financial situation:
    - Revenue: ₹{financial_data.get('revenue', 0):,.2f}
    - Profit: ₹{financial_data.get('profit', 0):,.2f}
    - Expenses: ₹{financial_data.get('expense', 0):,.2f}
    - Cash Balance: ₹{financial_data.get('cash_balance', 0):,.2f}
    - Profit Margin: {financial_data.get('profit_margin', 0):.1f}%
    - Runway: {financial_data.get('runway_months', 0):.1f} months
    
    User Query: {query}
    
    Provide specific, actionable financial advice based on their data. Be concise but helpful.
    """
    
    # For hackathon demo, you can use rule-based responses or integrate with OpenAI API
    return generate_rule_based_advice(query, financial_data)

def generate_rule_based_advice(query, data):
    """Rule-based financial advice for hackathon demo"""
    query_lower = query.lower()
    
    if 'cash flow' in query_lower or 'cashflow' in query_lower:
        runway = data.get('runway_months', 0)
        if runway < 3:
            return "⚠️ Critical: Your runway is less than 3 months. Consider: 1) Reducing non-essential expenses, 2) Accelerating revenue collection, 3) Exploring emergency funding options."
        elif runway < 6:
            return "⚡ Warning: 6-month runway approaching. Recommendations: 1) Review and optimize operational costs, 2) Focus on high-margin products/services, 3) Prepare contingency plans."
        else:
            return "✅ Healthy cash position! Consider: 1) Strategic investments for growth, 2) Building emergency reserves, 3) Optimizing cash deployment."
    
    elif 'profit' in query_lower or 'margin' in query_lower:
        margin = data.get('profit_margin', 0)
        if margin < 10:
            return "📈 Low profit margins detected. Strategies: 1) Analyze cost structure for optimization opportunities, 2) Review pricing strategy, 3) Focus on high-margin offerings."
        elif margin < 20:
            return "📊 Moderate margins. Growth opportunities: 1) Scale efficient operations, 2) Negotiate better supplier terms, 3) Automate repetitive processes."
        else:
            return "🎯 Strong margins! Consider: 1) Reinvesting in growth initiatives, 2) Market expansion, 3) Innovation and R&D investments."
    
    elif 'revenue' in query_lower or 'growth' in query_lower:
        revenue = data.get('revenue', 0)
        return f"💼 Current revenue: ₹{revenue:,.2f}. Growth strategies: 1) Customer acquisition campaigns, 2) Product/service diversification, 3) Market penetration analysis, 4) Strategic partnerships."
    
    elif 'expense' in query_lower or 'cost' in query_lower:
        expense = data.get('expense', 0)
        revenue = data.get('revenue', 1)
        expense_ratio = (expense / revenue) * 100 if revenue > 0 else 0
        return f"💰 Expenses are {expense_ratio:.1f}% of revenue. Optimization areas: 1) Fixed vs variable cost analysis, 2) Vendor negotiations, 3) Process automation, 4) Resource allocation review."
    
    else:
        return "🤖 I'm your CFO AI Assistant! I can help you with cash flow analysis, profit optimization, revenue growth strategies, and expense management. What specific financial area would you like to explore?"

@login_required(login_url='CFO:login')
def smart_insights(request):
    """Generate automated financial insights"""
    reports = FinancialReport.objects.filter(user=request.user).order_by('-uploaded_at')[:3]
    
    insights = []
    
    if reports.count() >= 2:
        current = reports[0]
        previous = reports[1]
        
        # Revenue trend analysis
        if current.total_revenue and previous.total_revenue:
            revenue_change = ((current.total_revenue - previous.total_revenue) / previous.total_revenue) * 100
            if revenue_change > 10:
                insights.append({
                    'type': 'positive',
                    'title': 'Strong Revenue Growth',
                    'message': f'Revenue increased by {revenue_change:.1f}% from last period.',
                    'action': 'Consider scaling successful initiatives.'
                })
            elif revenue_change < -10:
                insights.append({
                    'type': 'warning',
                    'title': 'Revenue Decline',
                    'message': f'Revenue decreased by {abs(revenue_change):.1f}% from last period.',
                    'action': 'Analyze market conditions and customer feedback.'
                })
        
        # Cash flow analysis
        if current.runway_months and current.runway_months < 6:
            insights.append({
                'type': 'critical',
                'title': 'Cash Flow Alert',
                'message': f'Only {current.runway_months:.1f} months of runway remaining.',
                'action': 'Implement cash conservation measures immediately.'
            })
        
        # Profit margin insights
        if current.profit_margin:
            if current.profit_margin > 25:
                insights.append({
                    'type': 'positive',
                    'title': 'Excellent Margins',
                    'message': f'Profit margin of {current.profit_margin:.1f}% is above industry average.',
                    'action': 'Maintain efficiency while exploring growth opportunities.'
                })
            elif current.profit_margin < 5:
                insights.append({
                    'type': 'warning',
                    'title': 'Thin Margins',
                    'message': f'Profit margin of {current.profit_margin:.1f}% needs improvement.',
                    'action': 'Review cost structure and pricing strategy.'
                })
    
    return render(request, 'CFO/smart_insights.html', {'insights': insights})

# ---------------- File Extraction ----------------
def extract_pdf_to_df(path):
    df = None
    try:
        with pdfplumber.open(path) as pdf:
            all_rows = []
            all_text = ""
            for page in pdf.pages:
                table = page.extract_table()
                text = page.extract_text() or ""
                all_text += text + "\n"
                if table:
                    header = table[0][:3]
                    data = [row[:3] for row in table[1:]]
                    all_rows.extend(data)

            if all_rows:
                df = pd.DataFrame(all_rows, columns=['Month','Account','Amount'])
            else:
                # Regex fallback
                patterns = {"Revenue": r"revenue.*?([\d,]+)",
                            "Net Income": r"net\s*income.*?([\d,]+)",
                            "Cash": r"cash.*?([\d,]+)"}
                numbers = {}
                for label, pattern in patterns.items():
                    match = re.search(pattern, all_text, re.I)
                    if match:
                        numbers[label] = float(match.group(1).replace(",", ""))
                if numbers:
                    df = pd.DataFrame({
                        'Month': ['Full Year'] * len(numbers),
                        'Account': list(numbers.keys()),
                        'Amount': list(numbers.values())
                    })
    except Exception as e:
        print("Error reading PDF:", e)
    return df


def extract_docx_to_df(path):
    df = None
    try:
        doc = Document(path)
        all_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                all_rows.append(cells)

        if all_rows:
            max_cols = max(len(r) for r in all_rows)
            cols = [f"col{i}" for i in range(max_cols)]
            df = pd.DataFrame([r + [""]*(max_cols-len(r)) for r in all_rows], columns=cols)
            if df.shape[1] >= 3:
                try:
                    df["Month"] = pd.to_datetime(df[df.columns[0]], errors="coerce").dt.to_period("M").astype(str)
                    df["Account"] = df[df.columns[1]]
                    df["Amount"] = pd.to_numeric(df[df.columns[2]].replace(r"[^0-9.\-]", "", regex=True), errors="coerce").fillna(0)
                except:
                    df["Month"] = "Full Year"
                    df["Account"] = df[df.columns[1]]
                    df["Amount"] = pd.to_numeric(df[df.columns[2]].replace(r"[^0-9.\-]", "", regex=True), errors="coerce").fillna(0)
            elif df.shape[1] == 2:
                df["Month"] = "Full Year"
                df["Account"] = df[df.columns[0]]
                df["Amount"] = pd.to_numeric(df[df.columns[1]].replace(r"[^0-9.\-]", "", regex=True), errors="coerce").fillna(0)
            else:
                df["Month"] = "Full Year"
                df["Account"] = "Unknown"
                df["Amount"] = 0
            return df

        # Paragraph regex fallback
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        pattern = re.compile(r"(?i)(revenue|net\s*income|net\s*profit|cash(?:\s*at\s*year-?end)?)[^:\n\r][:\s]\$?\s*([\d,]+(?:\.\d+)?)")
        matches = pattern.findall(full_text)
        if matches:
            numbers = {label.title(): float(num.replace(",", "")) for label, num in matches}
            df = pd.DataFrame({
                "Month": ["Full Year"] * len(numbers),
                "Account": list(numbers.keys()),
                "Amount": list(numbers.values())
            })
    except Exception as e:
        print("Error reading DOCX:", e)
    return df


def extract_image_to_df(path):
    df = None
    try:
        text = pytesseract.image_to_string(path)
        lines = []
        for line in text.split("\n"):
            parts = line.split()
            if len(parts) >= 2:
                lines.append(parts[:3])
        if lines:
            df = pd.DataFrame(lines, columns=['Month','Account','Amount'])
    except Exception as e:
        print("Error reading image:", e)
    return df


# ---------------- KPI Computation ----------------
def compute_kpis(report):
    try:
        path = report.file.path
        ext = os.path.splitext(path)[1].lower()
        df = None

        if ext in ['.xlsx', '.xls']:
            df = pd.read_excel(path)
        elif ext == '.csv':
            df = pd.read_csv(path)
        elif ext == '.json':
            df = pd.read_json(path)
        elif ext == '.pdf':
            df = extract_pdf_to_df(path)
        elif ext == '.docx':
            df = extract_docx_to_df(path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            df = extract_image_to_df(path)

        df = normalize_financial_df(df)

        if df.empty:
            return report

        revenue_mask = df['Account'].str.contains(r"rev|sales|income", case=False, na=False)
        expense_mask = df['Account'].str.contains(r"exp|cogs|payroll|cost|debit", case=False, na=False)

        revenue = df.loc[revenue_mask, "Amount"].sum()
        expense = df.loc[expense_mask, "Amount"].sum()

        profit = revenue - expense
        cash_balance = profit

        report.total_revenue = revenue
        report.total_expense = expense
        report.total_profit = profit
        report.cash_balance = cash_balance

        monthly_rev = df.loc[revenue_mask].groupby("Month")["Amount"].sum()
        if len(monthly_rev) > 1:
            prev, curr = monthly_rev.iloc[-2], monthly_rev.iloc[-1]
            report.revenue_growth = ((curr - prev) / max(prev, 1)) * 100
        else:
            report.revenue_growth = 0

        report.profit_margin = (profit / revenue * 100) if revenue else 0
        report.burn_rate = expense / max(len(df["Month"].unique()), 1)
        report.runway_months = cash_balance / max(report.burn_rate, 1)

        report.risk_flag = (
            "High Risk" if report.runway_months < 3
            else "Moderate/OK" if report.runway_months < 6
            else "Low Risk"
        )

        report.save()
    except Exception as e:
        print("Error computing KPIs:", e)

    return report


# ---------------- Views ----------------
@login_required(login_url='CFO:login')
@never_cache
def upload_report(request):
    if request.method == "POST":
        form = FinancialReportForm(request.POST, request.FILES)
        if form.is_valid():
            report = form.save(commit=False)
            report.user = request.user
            if not report.title:
                report.title = os.path.basename(report.file.name)
            report.save()
            report = compute_kpis(report)
            messages.success(request, "Report uploaded and KPIs computed!")
            return redirect("CFO:report_detail", pk=report.pk)
    else:
        form = FinancialReportForm()

    recent_reports = FinancialReport.objects.filter(user=request.user).order_by("-uploaded_at")[:5]
    return render(request, "CFO/upload_report.html", {"form": form, "recent_reports": recent_reports})


@login_required(login_url='CFO:login')
@never_cache
def report_detail(request, pk):
    report = get_object_or_404(FinancialReport, id=pk)
    path = report.file.path
    ext = os.path.splitext(path)[1].lower()
    df = None

    try:
        if ext in ['.xlsx', '.xls']:
            df = pd.read_excel(path)
        elif ext == '.csv':
            df = pd.read_csv(path)
        elif ext == '.json':
            df = pd.read_json(path)
        elif ext == '.pdf':
            df = extract_pdf_to_df(path)
        elif ext == '.docx':
            df = extract_docx_to_df(path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            df = extract_image_to_df(path)
    except Exception as e:
        print("Error reading file:", e)
        df = pd.DataFrame()

    df = normalize_financial_df(df)

    chart_data = {'labels': [], 'revenue': [], 'profit': [], 'expense': [], 'cash': []}
    if not df.empty:
        months = df['Month'].unique().tolist()
        chart_data['labels'] = months
        for m in months:
            month_df = df[df['Month'] == m]
            rev = month_df[month_df['Account'].str.contains(r"rev|sales|income", case=False, na=False)]["Amount"].sum()
            exp = month_df[month_df['Account'].str.contains(r"exp|cogs|payroll|cost|debit", case=False, na=False)]["Amount"].sum()
            profit = rev - exp
            cash = profit
            chart_data['revenue'].append(float(rev))
            chart_data['expense'].append(float(exp))
            chart_data['profit'].append(float(profit))
            chart_data['cash'].append(float(cash))

    total_revenue = float(sum(chart_data['revenue']))
    total_expense = float(sum(chart_data['expense']))
    total_profit = float(sum(chart_data['profit']))
    cash_balance = float(sum(chart_data['cash']))
    revenue_growth = ((chart_data['revenue'][-1] - chart_data['revenue'][-2]) / chart_data['revenue'][-2] * 100) \
        if len(chart_data['revenue']) >= 2 and chart_data['revenue'][-2] != 0 else 0
    profit_margin = (total_profit / total_revenue * 100) if total_revenue else 0
    burn_rate = total_expense / max(len(chart_data['labels']), 1)
    runway = cash_balance / burn_rate if burn_rate else 0

    kpis = {
        'total_revenue': total_revenue,
        'total_expense': total_expense,
        'total_profit': total_profit,
        'cash_balance': cash_balance,
        'revenue_growth': revenue_growth,
        'profit_margin': profit_margin,
        'burn_rate': burn_rate,
        'runway': runway,
    }

    return render(request, "CFO/report_detail.html", {
        "report": report,
        "kpis": kpis,
        "chart_data": json.loads(json.dumps(chart_data))
    })


@login_required(login_url='CFO:login')
@never_cache
def dashboard(request):
    reports = FinancialReport.objects.filter(user=request.user).order_by("-uploaded_at")
    return render(request, "CFO/dashboard.html", {"reports": reports})


# ---------------- Auth ----------------
@never_cache
def signup_view(request):
    if request.method == "POST":
        form = SignUpForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Account created successfully! Please log in.")
            return redirect('CFO:login')
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        form = SignUpForm()
    return render(request, 'CFO/signup.html', {'form': form})


@never_cache
def login_view(request):
    if request.method == "POST":
        form = LoginForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('CFO:dashboard')
        else:
            messages.error(request, "Invalid username or password.")
    else:
        form = LoginForm()
    return render(request, 'CFO/login.html', {'form': form})


@login_required(login_url='CFO:login')
def logout_view(request):
    logout(request)
    messages.info(request, "Logged out successfully")
    return redirect('CFO:login')


# ---------------- Landing ----------------
def index(request):
    return render(request, 'CFO/landing.html')


# ---------------- Scenario Planning ----------------
@login_required(login_url='CFO:login')
def scenario_planning(request):
    form = ScenarioForm(request.POST or None)
    results = []
    chart_data = None

    # Default monthly values
    base_monthly_revenue = 20000.0
    base_monthly_expense = 14000.0
    base_cash = 50000.0

    # Load latest financial report for base values
    latest = FinancialReport.objects.order_by('-uploaded_at').first()
    if latest:
        try:
            base_monthly_revenue = float(latest.total_revenue or base_monthly_revenue) / 12.0
            base_monthly_expense = float(latest.total_expense or base_monthly_expense) / 12.0
            base_cash = float(latest.cash_balance or base_cash)
        except Exception:
            pass

    if request.method == "POST" and form.is_valid():
        rev_growth_pct = form.cleaned_data.get('revenue_growth') or 0.0
        exp_growth_pct = form.cleaned_data.get('expense_growth') or 0.0
        cash_infusion = form.cleaned_data.get('cash_infusion') or 0.0
        months = form.cleaned_data.get('months') or 6

        rev_growth = float(rev_growth_pct) / 100.0
        exp_growth = float(exp_growth_pct) / 100.0

        revenue = float(base_monthly_revenue)
        expense = float(base_monthly_expense)
        cash = float(base_cash) + float(cash_infusion)

        for m in range(1, months + 1):
            revenue = revenue * (1 + rev_growth)
            expense = expense * (1 + exp_growth)
            cash = cash + (revenue - expense)
            results.append({
                "Month": f"Month {m}",
                "Revenue": round(revenue, 2),
                "Expense": round(expense, 2),
                "Cash": round(cash, 2)
            })

        chart_data = json.dumps(results)

    projected_revenue = sum(r['Revenue'] for r in results) if results else 0
    projected_expense = sum(r['Expense'] for r in results) if results else 0
    projected_profit = projected_revenue - projected_expense
    projected_cash = results[-1]['Cash'] if results else base_cash

    context = {
        "form": form,
        "results": results,
        "chart_data": chart_data,
        "projected_revenue": projected_revenue,
        "projected_expense": projected_expense,
        "projected_profit": projected_profit,
        "projected_cash": projected_cash,
    }
    return render(request, "CFO/scenario_planning.html", context)


# ---------------- Goal Tracking ----------------
@login_required(login_url='CFO:login')
def goal_tracking(request):
    form = GoalForm(request.POST or None)
    goals = FinancialGoal.objects.filter(user=request.user).order_by("-created_at")
    if request.method == "POST" and form.is_valid():
        goal = form.save(commit=False)
        goal.user = request.user
        goal.save()
        messages.success(request, "Goal added successfully!")
        return redirect("CFO:goal_tracking")

    latest_report = FinancialReport.objects.order_by("-uploaded_at").first()

    mapping = {"revenue": "total_revenue", "profit": "total_profit", "expense": "total_expense", "cash": "cash_balance"}
    progress_items = []
    for goal in goals:
        field_name = mapping.get(goal.goal_type)
        actual = getattr(latest_report, field_name) if latest_report and field_name else 0.0
        target = float(goal.target_value or 0.0)
        pct_rounded = round((actual / target * 100.0) if target else 0.0, 2)
        progress_width = max(0.0, min(pct_rounded, 100.0))
        status = "On Track ✅" if pct_rounded >= 100.0 else "In Progress ⏳"

        progress_items.append({
            "goal": goal,
            "actual": actual,
            "progress_pct": pct_rounded,
            "progress_width": progress_width,
            "status": status,
        })

    return render(request, "CFO/goal_tracking.html", {"form": form, "progress_items": progress_items})


# ---------------- Export Center ----------------
@login_required(login_url='CFO:login')
def export_center(request):
    return render(request, "CFO/export_center.html")


@login_required(login_url='CFO:login')
def export_reports_csv(request):
    reports = FinancialReport.objects.all()
    response = HttpResponse(content_type="text/csv")
    response['Content-Disposition'] = 'attachment; filename="financial_reports.csv"'
    writer = csv.writer(response)
    writer.writerow(["Revenue", "Profit", "Expense", "Cash Balance", "Uploaded At"])
    for r in reports:
        writer.writerow([r.total_revenue, r.total_profit, r.total_expense, r.cash_balance, r.uploaded_at])
    return response


@login_required(login_url='CFO:login')
def export_reports_excel(request):
    reports = FinancialReport.objects.all()
    wb = Workbook()
    ws = wb.active
    ws.append(["Revenue", "Profit", "Expense", "Cash Balance", "Uploaded At"])
    for r in reports:
        ws.append([r.total_revenue, r.total_profit, r.total_expense, r.cash_balance, str(r.uploaded_at)])
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response['Content-Disposition'] = 'attachment; filename="financial_reports.xlsx"'
    wb.save(response)
    return response


@login_required(login_url='CFO:login')
def export_reports_pdf(request):
    reports = FinancialReport.objects.all()
    response = HttpResponse(content_type="application/pdf")
    response['Content-Disposition'] = 'attachment; filename="financial_reports.pdf"'
    p = canvas.Canvas(response)
    p.setFont("Helvetica", 14)
    p.drawString(100, 800, "Financial Reports Export")
    y = 760
    for r in reports:
        line = f"Revenue: {r.total_revenue}, Profit: {r.total_profit}, Expense: {r.total_expense}, Cash: {r.cash_balance}, Uploaded: {r.uploaded_at}"
        p.drawString(80, y, line)
        y -= 20
    p.showPage()
    p.save()
    return response


# ---------------- Financial Analysis ----------------
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import FinancialReport
from django.utils.safestring import mark_safe

@login_required(login_url='CFO:login')
def financial_analysis(request):
    reports = FinancialReport.objects.all().order_by("uploaded_at")
    if not reports.exists():
        return render(request, "CFO/financial_analysis.html", {"error": "No reports uploaded yet."})

    # Extract numeric data safely
    revenues = [float(r.total_revenue or 0) for r in reports]
    profits = [float(r.total_profit or 0) for r in reports]
    expenses = [float(r.total_expense or 0) for r in reports]
    cash_balances = [float(r.cash_balance or 0) for r in reports]
    dates = [r.uploaded_at.strftime("%b %Y") for r in reports]  # e.g., "Jan 2025"

    # Latest report for summary cards
    latest = reports.last()
    total_revenue = float(latest.total_revenue or 0)
    total_profit = float(latest.total_profit or 0)
    total_expense = float(latest.total_expense or 0)
    cash_balance = float(latest.cash_balance or 0)

    # Key metrics
    profit_margin = round((total_profit / total_revenue) * 100, 2) if total_revenue else 0
    burn_rate = round(total_expense / 12, 2) if total_expense else 0
    runway = round(cash_balance / burn_rate, 1) if burn_rate else float("inf")
    growth = round(((revenues[-1] - revenues[0]) / revenues[0]) * 100, 2) if revenues and revenues[0] else 0

    context = {
        # Chart.js data (dump to JSON to avoid escaping issues)
        "months": mark_safe(json.dumps(dates)),
        "profit_trend": mark_safe(json.dumps(profits)),
        "revenue_trend": mark_safe(json.dumps(revenues)),
        "expense_trend": mark_safe(json.dumps(expenses)),
        "cash_trend": mark_safe(json.dumps(cash_balances)),

        # Summary cards
        "total_revenue": total_revenue,
        "total_expenses": total_expense,
        "net_profit": total_profit,
        "profit_margin": profit_margin,
        "burn_rate": burn_rate,
        "runway": runway,
        "growth": growth,
    }
    return render(request, "CFO/financial_analysis.html", context)



# ---------------- Cash Flow View ----------------
@login_required(login_url='CFO:login')
def cash_flow_view(request):
    # Sample data
    cash_inflows = 500000
    cash_outflows = 300000
    net_cash_flow = cash_inflows - cash_outflows

    months = ["Jan", "Feb", "Mar", "Apr", "May"]
    net_cash_flow_trend = [50000, 60000, 55000, 70000, 65000]

    context = {
        "cash_inflows": cash_inflows,
        "cash_outflows": cash_outflows,
        "net_cash_flow": net_cash_flow,
        # JSON-safe for JS
        "months": mark_safe(json.dumps(months)),
        "net_cash_flow_trend": mark_safe(json.dumps(net_cash_flow_trend)),
    }
    return render(request, "CFO/cash_flow.html", context)


import json
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .forms import ReportFilterForm
from .models import FinancialReport

@login_required(login_url='CFO:login')
def reports_generator(request):
    form = ReportFilterForm(request.POST or None)
    report = None
    chart_data = None
    report_type = None

    if request.method == 'POST' and form.is_valid():
        start_date = form.cleaned_data['start_date']
        end_date = form.cleaned_data['end_date']
        report_type = form.cleaned_data['report_type']  # e.g., 'Monthly', 'Quarterly'

        # Filter reports by date range
        reports_qs = FinancialReport.objects.filter(
            uploaded_at__date__range=(start_date, end_date)
        ).order_by('uploaded_at')

        if reports_qs.exists():
            # Pick the latest report for the KPIs
            report = reports_qs.last()

            # Auto-calculate missing fields if not present
            if report.total_profit is None:
                report.total_profit = (report.total_revenue or 0) - (report.total_expense or 0)
            if report.profit_margin is None:
                report.profit_margin = ((report.total_profit or 0) / max(report.total_revenue or 1, 1)) * 100
            if report.burn_rate is None:
                report.burn_rate = (report.total_expense or 0)
            if report.runway_months is None:
                report.runway_months = ((report.cash_balance or 0) / max(report.burn_rate or 1, 1))

            # Prepare chart data (last 12 reports or all)
            chart_data = []
            for r in reports_qs[:12]:
                chart_data.append({
                    "Month": r.uploaded_at.strftime('%b %Y'),
                    "Revenue": r.total_revenue or 0,
                    "Profit": (r.total_profit or (r.total_revenue or 0) - (r.total_expense or 0)),
                    "Cash": r.cash_balance or 0
                })
            chart_data = json.dumps(chart_data)
        else:
            report = None
            chart_data = None

    context = {
        'form': form,
        'report': report,
        'chart_data': chart_data,
        'report_type': report_type
    }
    return render(request, 'CFO/reports_generator.html', context)


@login_required(login_url='CFO:login')
def performance_metrics(request):
    # Fetch the latest 12 months (or all) reports to calculate trends
    reports = FinancialReport.objects.order_by("-uploaded_at")[:12][::-1]  # Oldest first

    kpis = {}
    if reports:
        latest = reports[-1]

        # Total Revenue, Profit, Expense, Cash
        total_revenue = latest.total_revenue or 0.0
        total_expense = latest.total_expense or 0.0
        total_profit = (total_revenue - total_expense)
        cash_balance = latest.cash_balance or 0.0

        # Burn Rate: average monthly expense over available reports
        burn_rate = sum(r.total_expense or 0.0 for r in reports) / max(len(reports), 1)

        # Runway: months until cash runs out
        runway_months = cash_balance / max(burn_rate, 1)

        # Profit Margin
        profit_margin = (total_profit / max(total_revenue, 1)) * 100

        # Revenue Growth: compare first and last in period
        first_revenue = reports[0].total_revenue or 1
        revenue_growth = ((total_revenue - first_revenue) / first_revenue) * 100

        # Populate KPI dict
        kpis = {
            "total_revenue": round(total_revenue, 2),
            "total_profit": round(total_profit, 2),
            "total_expense": round(total_expense, 2),
            "cash_balance": round(cash_balance, 2),
            "burn_rate": round(burn_rate, 2),
            "runway_months": round(runway_months, 1),
            "profit_margin": round(profit_margin, 2),
            "revenue_growth": round(revenue_growth, 2),
            # Percentages for progress bars
            "profit_margin_pct": min(max(profit_margin, 0), 100),
            "revenue_growth_pct": min(max(revenue_growth, 0), 100),
            "burn_rate_pct": min(max((burn_rate / max(total_revenue, 1)) * 100, 0), 100)
        }
    else:
        # Default empty KPIs
        kpis = {
            "total_revenue": 0,
            "total_profit": 0,
            "total_expense": 0,
            "cash_balance": 0,
            "burn_rate": 0,
            "runway_months": 0,
            "profit_margin": 0,
            "revenue_growth": 0,
            "profit_margin_pct": 0,
            "revenue_growth_pct": 0,
            "burn_rate_pct": 0
        }

    return render(request, "CFO/performance_metrics.html", {"kpis": kpis})


#---------Setting Page -----------
from django.contrib.auth import update_session_auth_hash
@login_required
def settings_view(request):
    user = request.user
    profile, created = Profile.objects.get_or_create(user=user)

    if request.method == 'POST':
        form_type = request.POST.get('form_type')

        if form_type == 'profile':
            profile_form = ProfileForm(request.POST, instance=user)
            if profile_form.is_valid():
                # Ensure email does not get overwritten
                profile_form.cleaned_data.pop('email', None)
                profile_form.save()
                messages.success(request, "Profile updated successfully!")
                return redirect('CFO:settings_view')
            else:
                messages.error(request, "Please fix errors in profile form.")

        elif form_type == 'preferences':
            pref_form = PreferencesForm(request.POST, instance=profile)
            if pref_form.is_valid():
                pref_form.save()
                messages.success(request, "Preferences updated successfully!")
                return redirect('CFO:settings_view')
            else:
                messages.error(request, "Please fix errors in preferences form.")

        elif form_type == 'password':
            pwd_form = SetPasswordForm(user=user, data=request.POST)
            if request.POST.get("current_password") and user.check_password(request.POST.get("current_password")):
                if pwd_form.is_valid():
                    pwd_form.save()
                    update_session_auth_hash(request, user)
                    messages.success(request, "Password changed successfully!")
                    return redirect('CFO:settings_view')
            else:
                messages.error(request, "Current password is incorrect.")
    else:
        profile_form = ProfileForm(instance=user)
        pref_form = PreferencesForm(instance=profile)
        pwd_form = SetPasswordForm(user=user)

    return render(request, 'CFO/settings.html', {
        'profile_form': profile_form,
        'pref_form': pref_form,
        'pwd_form': pwd_form,
    })

#----help page----
def help_view(request):
    return render(request, 'CFO/help.html')

#----contact page----
def contact_view(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        message_text = request.POST.get('message')

        # Save to database
        contact = Contact(name=name, email=email, message=message_text)
        contact.save()

        # Add a success message
        messages.success(request, 'Your message has been sent!')

        return redirect('CFO:contact_view')  # redirect to clear POST data

    return render(request, 'CFO/contact.html')




#----delete account----
@login_required
def delete_account(request):
    """
    Delete the currently-logged-in user's account.
    Only accepts POST to avoid CSRF-less deletes.
    Logs the user out then deletes the User record and redirects to home.
    """
    if request.method == "POST":
        user = request.user
        # optional: store username/email for logging before deletion
        logout(request)
        user.delete()
        messages.success(request, "Your account has been deleted.")
        return redirect("/")   # or change to a named route like 'index' if you have one

    # If someone visits the URL via GET, just redirect back to settings
    messages.error(request, "Invalid request method.")
    return redirect("CFO:settings_view")




@login_required(login_url='CFO:login')
def download_report_pdf(request, pk):
    report = get_object_or_404(FinancialReport, id=pk)

    # KPIs
    kpis = {
        "total_revenue": getattr(report, 'total_revenue', 0),
        "total_profit": getattr(report, 'total_profit', 0),
        "total_expense": getattr(report, 'total_expense', 0),
        "cash_balance": getattr(report, 'cash_balance', 0),
        "revenue_growth": getattr(report, 'revenue_growth', None),
        "profit_margin": getattr(report, 'profit_margin', None),
        "runway": getattr(report, 'runway', None),
        "burn_rate": getattr(report, 'burn_rate', None)
    }

    # Chart images (all dynamically generated graphs)
    chart_images = {}
    if hasattr(report, 'charts') and isinstance(report.charts, dict):
        # report.charts assumed to be a dict of {chart_name: plt.Figure}
        def chart_to_base64(fig):
            buf = io.BytesIO()
            fig.savefig(buf, format='png', bbox_inches='tight')
            buf.seek(0)
            return "data:image/png;base64," + base64.b64encode(buf.read()).decode()

        for name, fig in report.charts.items():
            chart_images[name] = chart_to_base64(fig)
            plt.close(fig)

    # Render PDF
    html = render_to_string('CFO/report_detail_pdf.html', {
        'report': report,
        'kpis': kpis,
        'chart_images': chart_images
    })

    result = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html), dest=result)
    if pisa_status.err:
        return HttpResponse('Error generating PDF', status=500)

    response = HttpResponse(result.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="report_{report.id}.pdf"'
    return response


def financial_health(request):
    context = {
        "health_score": 82,
        "strengths": ["Good Revenue Growth", "Strong Cash Flow"],
        "weaknesses": ["High Debt Ratio", "Low Liquidity"],
    }
    return render(request, 'CFO/financial_health.html', context)


def risk_alerts(request):
    """Generate risk alerts from the latest uploaded report"""
    report = FinancialReport.objects.filter(user=request.user).order_by('-uploaded_at').first()
    
    alerts = []

    if report:
        # ✅ Example rules using your actual fields
        if report.profit_margin is not None and report.profit_margin < 5:
            alerts.append("🚨 Low Profit Margin: Less than 5%, sustainability risk.")

        if report.revenue_growth is not None and report.revenue_growth < 0:
            alerts.append("⚠ Negative Revenue Growth: Business is shrinking.")

        if report.burn_rate is not None and report.runway_months is not None:
            if report.runway_months < 6:
                alerts.append(f"🔥 High Burn Rate: Runway only {report.runway_months} months left.")

        if report.cash_balance is not None and report.cash_balance < 10000:
            alerts.append("💸 Low Cash Balance: Cash reserves critically low.")

        # Use risk_flag field if set
        if report.risk_flag:
            alerts.append(f"📊 Risk Classification: {report.risk_flag}")

        if not alerts:
            alerts.append("✅ No major financial risks detected in the latest report.")
    else:
        alerts.append("❌ No reports uploaded yet. Please upload to see risk alerts.")

    context = {"alerts": alerts, "report": report}
    return render(request, "CFO/risk_alerts.html", context)

#----profile page-----
@login_required
def profile_view(request):
    return render(request, "CFO/profile.html")